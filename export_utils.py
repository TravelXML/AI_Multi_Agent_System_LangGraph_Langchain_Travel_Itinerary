import io
import os

from markdown import markdown as md_to_html
from bs4 import BeautifulSoup, NavigableString, Tag
from xhtml2pdf import pisa
from docx import Document
from docx.shared import Pt

MD_EXTENSIONS = ["tables", "fenced_code", "sane_lists"]

_FONT_REG = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
_FONT_BOLD = "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"
_HAS_UNICODE_FONT = os.path.isfile(_FONT_REG) and os.path.isfile(_FONT_BOLD)

_FONT_FACE_CSS = f"""
@font-face {{ font-family: "DejaVu"; src: url("{_FONT_REG}"); }}
@font-face {{ font-family: "DejaVu"; font-weight: bold; src: url("{_FONT_BOLD}"); }}
""" if _HAS_UNICODE_FONT else ""

_BODY_FONT = '"DejaVu", sans-serif' if _HAS_UNICODE_FONT else "sans-serif"

PDF_CSS = f"""
{_FONT_FACE_CSS}
body {{ font-family: {_BODY_FONT}; font-size: 10pt; color: #1a1a1a; line-height: 1.5; }}
h1 {{ font-size: 20pt; color: #0d4a8a; margin-bottom: 4pt; }}
h2 {{ font-size: 15pt; color: #0d4a8a; margin-top: 16pt; border-bottom: 1pt solid #cccccc; padding-bottom: 3pt; }}
h3 {{ font-size: 12pt; color: #1a6bbf; margin-top: 12pt; }}
h4 {{ font-size: 11pt; color: #1a6bbf; margin-top: 10pt; }}
table {{ border-collapse: collapse; width: 100%; margin: 8pt 0; }}
th {{ background-color: #eaf1fb; border: 0.75pt solid #99b8d9; padding: 5pt 7pt; text-align: left; font-weight: bold; }}
td {{ border: 0.75pt solid #cccccc; padding: 5pt 7pt; }}
hr {{ border: none; border-top: 0.75pt solid #cccccc; margin: 10pt 0; }}
code {{ font-family: {_BODY_FONT}; background-color: #f2f2f2; padding: 1pt 3pt; }}
pre {{ background-color: #f2f2f2; padding: 6pt; }}
strong {{ font-weight: bold; }}
ul, ol {{ margin: 4pt 0 4pt 14pt; padding: 0; }}
li {{ margin-bottom: 2pt; }}
blockquote {{ color: #555555; border-left: 2pt solid #cccccc; padding-left: 8pt; margin-left: 0; }}
"""


def render_pdf(markdown_text: str) -> bytes:
    body_html = md_to_html(markdown_text, extensions=MD_EXTENSIONS)
    html = f"<html><head><style>{PDF_CSS}</style></head><body>{body_html}</body></html>"
    buf = io.BytesIO()
    pisa.CreatePDF(io.StringIO(html), dest=buf)
    return buf.getvalue()


def _add_inline(paragraph, node):
    for child in node.children:
        if isinstance(child, NavigableString):
            text = str(child)
            if text:
                paragraph.add_run(text)
        elif isinstance(child, Tag):
            if child.name in ("strong", "b"):
                paragraph.add_run(child.get_text()).bold = True
            elif child.name in ("em", "i"):
                paragraph.add_run(child.get_text()).italic = True
            elif child.name == "code":
                run = paragraph.add_run(child.get_text())
                run.font.name = "Consolas"
            elif child.name == "br":
                paragraph.add_run().add_break()
            else:
                _add_inline(paragraph, child)


def _set_table_style(table):
    for style_name in ("Light Grid Accent 1", "Table Grid"):
        try:
            table.style = style_name
            return
        except KeyError:
            continue


def render_docx(markdown_text: str) -> bytes:
    html = md_to_html(markdown_text, extensions=MD_EXTENSIONS)
    soup = BeautifulSoup(html, "html.parser")
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    for el in soup.contents:
        if not isinstance(el, Tag):
            continue

        if el.name in ("h1", "h2", "h3", "h4"):
            level = min(int(el.name[1]), 4)
            doc.add_heading(el.get_text(), level=level)

        elif el.name == "p":
            _add_inline(doc.add_paragraph(), el)

        elif el.name == "hr":
            doc.add_paragraph("─" * 40)

        elif el.name in ("ul", "ol"):
            style_name = "List Bullet" if el.name == "ul" else "List Number"
            for li in el.find_all("li", recursive=False):
                _add_inline(doc.add_paragraph(style=style_name), li)

        elif el.name == "table":
            rows = el.find_all("tr")
            if not rows:
                continue
            n_cols = max(len(r.find_all(["td", "th"])) for r in rows)
            table = doc.add_table(rows=0, cols=n_cols)
            _set_table_style(table)
            for r in rows:
                cells = r.find_all(["td", "th"])
                is_header = r.find("th") is not None
                row_cells = table.add_row().cells
                for i, c in enumerate(cells):
                    if i >= n_cols:
                        continue
                    para = row_cells[i].paragraphs[0]
                    _add_inline(para, c)
                    if is_header:
                        for run in para.runs:
                            run.bold = True

        elif el.name == "pre":
            para = doc.add_paragraph(el.get_text())
            for run in para.runs:
                run.font.name = "Consolas"
                run.font.size = Pt(9)

        elif el.name == "blockquote":
            para = doc.add_paragraph()
            _add_inline(para, el)
            for run in para.runs:
                run.italic = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
